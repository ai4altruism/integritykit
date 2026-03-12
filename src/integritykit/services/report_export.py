"""After-action report export service for Sprint 8.

Implements:
- S8-14: After-action report export (PDF/DOCX)
- Aggregates analytics data for post-incident analysis
- Generates structured reports with charts and tables
"""

import io
import logging
from datetime import datetime
from typing import Any

from integritykit.models.report import (
    AfterActionReportData,
    AfterActionReportRequest,
    CandidateSummary,
    ConflictSummary,
    FacilitatorSummary,
    ReportFormat,
    ReportSection,
    SignalSummary,
    TimelineEvent,
    TopicSummary,
)

logger = logging.getLogger(__name__)


class ReportExportService:
    """Service for generating after-action reports in PDF/DOCX format."""

    def __init__(
        self,
        signals_collection: Any = None,
        candidates_collection: Any = None,
        audit_log_collection: Any = None,
        clusters_collection: Any = None,
        users_collection: Any = None,
    ):
        """Initialize report export service.

        Args:
            signals_collection: MongoDB signals collection
            candidates_collection: MongoDB candidates collection
            audit_log_collection: MongoDB audit log collection
            clusters_collection: MongoDB clusters collection
            users_collection: MongoDB users collection
        """
        self.signals = signals_collection
        self.candidates = candidates_collection
        self.audit_log = audit_log_collection
        self.clusters = clusters_collection
        self.users = users_collection

    async def generate_report(
        self,
        request: AfterActionReportRequest,
    ) -> tuple[bytes, str]:
        """Generate after-action report.

        Args:
            request: Report request parameters

        Returns:
            Tuple of (report bytes, content type)

        Raises:
            ValueError: If request parameters are invalid
        """
        # Validate time range
        days = (request.end_date - request.start_date).days
        if days > 90:
            raise ValueError("Report time range cannot exceed 90 days")
        if days < 0:
            raise ValueError("End date must be after start date")

        # Aggregate report data
        report_data = await self._aggregate_report_data(request)

        # Generate report in requested format
        if request.format == ReportFormat.PDF:
            content = self._generate_pdf(report_data, request)
            content_type = "application/pdf"
        else:
            content = self._generate_docx(report_data, request)
            content_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        return content, content_type

    async def _aggregate_report_data(
        self,
        request: AfterActionReportRequest,
    ) -> AfterActionReportData:
        """Aggregate all data needed for the report.

        Args:
            request: Report request parameters

        Returns:
            Aggregated report data
        """
        report_data = AfterActionReportData(
            workspace_id=request.workspace_id,
            title=request.title,
            incident_name=request.incident_name,
            start_date=request.start_date,
            end_date=request.end_date,
        )

        # Aggregate sections based on request
        sections = request.sections

        if ReportSection.EXECUTIVE_SUMMARY in sections or ReportSection.SIGNAL_ANALYSIS in sections:
            report_data.signal_summary = await self._aggregate_signal_summary(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )
            report_data.candidate_summary = await self._aggregate_candidate_summary(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )

        if ReportSection.TIMELINE in sections:
            report_data.timeline = await self._build_timeline(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )

        if ReportSection.FACILITATOR_PERFORMANCE in sections:
            report_data.facilitator_summaries = await self._aggregate_facilitator_summaries(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )

        if ReportSection.CONFLICT_RESOLUTION in sections:
            report_data.conflict_summary = await self._aggregate_conflict_summary(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )

        if ReportSection.TOPIC_TRENDS in sections:
            report_data.topic_summaries = await self._aggregate_topic_summaries(
                request.workspace_id,
                request.start_date,
                request.end_date,
            )

        if ReportSection.RECOMMENDATIONS in sections:
            report_data.recommendations = self._generate_recommendations(report_data)

        return report_data

    async def _aggregate_signal_summary(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
    ) -> SignalSummary:
        """Aggregate signal summary statistics."""
        if self.signals is None:
            return SignalSummary(total_signals=0)

        # Aggregate signal volume by channel and time
        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "created_at": {"$gte": start_date, "$lte": end_date},
                }
            },
            {
                "$group": {
                    "_id": None,
                    "total": {"$sum": 1},
                    "by_channel": {
                        "$push": "$channel_id"
                    },
                }
            },
        ]

        cursor = self.signals.aggregate(pipeline)
        results = [doc async for doc in cursor]

        if not results:
            return SignalSummary(total_signals=0)

        result = results[0]
        total = result.get("total", 0)

        # Count by channel
        channels = result.get("by_channel", [])
        by_channel: dict[str, int] = {}
        for ch in channels:
            by_channel[ch] = by_channel.get(ch, 0) + 1

        # Calculate average per day
        days = max((end_date - start_date).days, 1)
        avg_per_day = total / days

        return SignalSummary(
            total_signals=total,
            signals_by_channel=by_channel,
            avg_signals_per_day=round(avg_per_day, 2),
        )

    async def _aggregate_candidate_summary(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
    ) -> CandidateSummary:
        """Aggregate candidate processing statistics."""
        if self.candidates is None:
            return CandidateSummary(total_candidates=0)

        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "created_at": {"$gte": start_date, "$lte": end_date},
                }
            },
            {
                "$group": {
                    "_id": "$readiness_state",
                    "count": {"$sum": 1},
                }
            },
        ]

        cursor = self.candidates.aggregate(pipeline)
        results = [doc async for doc in cursor]

        total = 0
        verified = 0
        blocked = 0
        in_review = 0

        for r in results:
            state = r["_id"]
            count = r["count"]
            total += count
            if state == "VERIFIED":
                verified = count
            elif state == "BLOCKED":
                blocked = count
            elif state == "IN_REVIEW":
                in_review = count

        verification_rate = verified / total if total > 0 else 0.0

        return CandidateSummary(
            total_candidates=total,
            verified_count=verified,
            blocked_count=blocked,
            in_review_count=in_review,
            verification_rate=round(verification_rate, 3),
        )

    async def _build_timeline(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
    ) -> list[TimelineEvent]:
        """Build timeline of significant events."""
        events: list[TimelineEvent] = []

        if self.audit_log is None:
            return events

        # Get significant audit events
        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "timestamp": {"$gte": start_date, "$lte": end_date},
                    "action_type": {
                        "$in": [
                            "cop_candidate.verify",
                            "cop_candidate.block",
                            "conflict.detected",
                            "conflict.resolved",
                            "cop_update.publish",
                        ]
                    },
                }
            },
            {"$sort": {"timestamp": 1}},
            {"$limit": 100},
        ]

        cursor = self.audit_log.aggregate(pipeline)

        action_descriptions = {
            "cop_candidate.verify": "Candidate verified",
            "cop_candidate.block": "Candidate blocked",
            "conflict.detected": "Conflict detected",
            "conflict.resolved": "Conflict resolved",
            "cop_update.publish": "COP update published",
        }

        action_significance = {
            "cop_update.publish": "notable",
            "conflict.detected": "notable",
            "cop_candidate.block": "notable",
        }

        async for doc in cursor:
            action = doc.get("action_type", "unknown")
            events.append(
                TimelineEvent(
                    timestamp=doc.get("timestamp", datetime.utcnow()),
                    event_type=action,
                    description=action_descriptions.get(action, action),
                    significance=action_significance.get(action, "normal"),
                    related_ids=[str(doc.get("target_id", ""))],
                )
            )

        return events

    async def _aggregate_facilitator_summaries(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
    ) -> list[FacilitatorSummary]:
        """Aggregate facilitator performance metrics."""
        summaries: list[FacilitatorSummary] = []

        if self.audit_log is None:
            return summaries

        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "timestamp": {"$gte": start_date, "$lte": end_date},
                    "actor_type": "facilitator",
                }
            },
            {
                "$group": {
                    "_id": "$actor_id",
                    "total_actions": {"$sum": 1},
                    "candidates": {"$addToSet": "$target_id"},
                    "action_types": {"$push": "$action_type"},
                }
            },
            {"$sort": {"total_actions": -1}},
            {"$limit": 20},
        ]

        cursor = self.audit_log.aggregate(pipeline)

        async for doc in cursor:
            user_id = doc["_id"]

            # Get user name if available
            user_name = None
            if self.users:
                user = await self.users.find_one({"_id": user_id})
                if user:
                    user_name = user.get("name") or user.get("display_name")

            action_types = doc.get("action_types", [])
            verifications = sum(1 for a in action_types if "verify" in a)
            resolutions = sum(1 for a in action_types if "resolve" in a)

            summaries.append(
                FacilitatorSummary(
                    user_id=user_id,
                    user_name=user_name,
                    total_actions=doc.get("total_actions", 0),
                    candidates_processed=len(doc.get("candidates", [])),
                    verification_actions=verifications,
                    conflict_resolutions=resolutions,
                )
            )

        return summaries

    async def _aggregate_conflict_summary(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
    ) -> ConflictSummary:
        """Aggregate conflict resolution metrics."""
        if self.clusters is None:
            return ConflictSummary()

        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "has_conflict": True,
                    "created_at": {"$gte": start_date, "$lte": end_date},
                }
            },
            {
                "$group": {
                    "_id": "$risk_tier",
                    "total": {"$sum": 1},
                    "resolved": {
                        "$sum": {"$cond": [{"$eq": ["$conflict_resolved", True]}, 1, 0]}
                    },
                    "resolution_methods": {"$push": "$resolution_method"},
                }
            },
        ]

        cursor = self.clusters.aggregate(pipeline)

        total_conflicts = 0
        resolved_conflicts = 0
        by_risk_tier: dict[str, int] = {}
        by_method: dict[str, int] = {}

        async for doc in cursor:
            tier = doc["_id"] or "unknown"
            tier_total = doc.get("total", 0)
            tier_resolved = doc.get("resolved", 0)

            total_conflicts += tier_total
            resolved_conflicts += tier_resolved
            by_risk_tier[tier] = tier_total

            for method in doc.get("resolution_methods", []):
                if method:
                    by_method[method] = by_method.get(method, 0) + 1

        resolution_rate = resolved_conflicts / total_conflicts if total_conflicts > 0 else 0.0

        return ConflictSummary(
            total_conflicts=total_conflicts,
            resolved_conflicts=resolved_conflicts,
            resolution_rate=round(resolution_rate, 3),
            by_risk_tier=by_risk_tier,
            by_resolution_method=by_method,
        )

    async def _aggregate_topic_summaries(
        self,
        workspace_id: str,
        start_date: datetime,
        end_date: datetime,
        limit: int = 10,
    ) -> list[TopicSummary]:
        """Aggregate top trending topics."""
        topics: list[TopicSummary] = []

        if self.signals is None:
            return topics

        pipeline = [
            {
                "$match": {
                    "workspace_id": workspace_id,
                    "created_at": {"$gte": start_date, "$lte": end_date},
                    "cluster_id": {"$exists": True},
                }
            },
            {
                "$lookup": {
                    "from": "clusters",
                    "localField": "cluster_id",
                    "foreignField": "_id",
                    "as": "cluster",
                }
            },
            {"$unwind": "$cluster"},
            {
                "$group": {
                    "_id": {
                        "cluster_id": "$cluster_id",
                        "topic": "$cluster.topic",
                        "topic_type": "$cluster.topic_type",
                    },
                    "signal_count": {"$sum": 1},
                    "first_seen": {"$min": "$created_at"},
                    "last_seen": {"$max": "$created_at"},
                }
            },
            {"$sort": {"signal_count": -1}},
            {"$limit": limit},
        ]

        cursor = self.signals.aggregate(pipeline)

        async for doc in cursor:
            id_info = doc["_id"]
            topics.append(
                TopicSummary(
                    topic=id_info.get("topic", "Unknown"),
                    topic_type=id_info.get("topic_type", "unknown"),
                    signal_count=doc.get("signal_count", 0),
                    first_seen=doc.get("first_seen"),
                    peak_time=doc.get("last_seen"),
                )
            )

        return topics

    def _generate_recommendations(
        self,
        report_data: AfterActionReportData,
    ) -> list[str]:
        """Generate recommendations based on report metrics."""
        recommendations: list[str] = []

        # Signal volume recommendations
        if report_data.signal_summary and report_data.signal_summary.total_signals == 0:
            recommendations.append(
                "No signals were ingested during this period. "
                "Verify that signal sources are properly configured."
            )

        # Candidate processing recommendations
        if report_data.candidate_summary:
            cs = report_data.candidate_summary
            if cs.verification_rate < 0.5 and cs.total_candidates > 10:
                recommendations.append(
                    f"Verification rate ({cs.verification_rate:.1%}) is below 50%. "
                    "Consider additional facilitator training or reviewing verification criteria."
                )
            if cs.blocked_count > cs.verified_count:
                recommendations.append(
                    "More candidates were blocked than verified. "
                    "Review conflict resolution procedures and escalation paths."
                )

        # Conflict resolution recommendations
        if report_data.conflict_summary:
            cf = report_data.conflict_summary
            if cf.resolution_rate < 0.7 and cf.total_conflicts > 5:
                recommendations.append(
                    f"Conflict resolution rate ({cf.resolution_rate:.1%}) is below 70%. "
                    "Consider streamlining resolution workflows or adding facilitator capacity."
                )

        # Facilitator workload recommendations
        if report_data.facilitator_summaries:
            actions = [f.total_actions for f in report_data.facilitator_summaries]
            if actions:
                max_actions = max(actions)
                min_actions = min(actions)
                if max_actions > 3 * min_actions and len(actions) > 1:
                    recommendations.append(
                        "Facilitator workload is unevenly distributed. "
                        "Consider load balancing or capacity planning."
                    )

        if not recommendations:
            recommendations.append(
                "Operations during this period were within normal parameters. "
                "Continue current procedures and monitoring."
            )

        return recommendations

    def _generate_pdf(
        self,
        report_data: AfterActionReportData,
        request: AfterActionReportRequest,
    ) -> bytes:
        """Generate PDF report using ReportLab."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story: list[Any] = []

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            spaceAfter=30,
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
        )
        body_style = styles["Normal"]

        # Title
        story.append(Paragraph(report_data.title, title_style))
        if report_data.incident_name:
            story.append(Paragraph(f"Incident: {report_data.incident_name}", body_style))
        story.append(
            Paragraph(
                f"Period: {report_data.start_date.strftime('%Y-%m-%d')} to "
                f"{report_data.end_date.strftime('%Y-%m-%d')}",
                body_style,
            )
        )
        story.append(
            Paragraph(
                f"Generated: {report_data.generated_at.strftime('%Y-%m-%d %H:%M UTC')}",
                body_style,
            )
        )
        story.append(Spacer(1, 0.3 * inch))

        # Executive Summary
        if ReportSection.EXECUTIVE_SUMMARY in request.sections:
            story.append(Paragraph("Executive Summary", heading_style))

            summary_data = []
            if report_data.signal_summary:
                summary_data.append(
                    ["Total Signals", str(report_data.signal_summary.total_signals)]
                )
                summary_data.append(
                    [
                        "Avg Signals/Day",
                        f"{report_data.signal_summary.avg_signals_per_day:.1f}",
                    ]
                )
            if report_data.candidate_summary:
                summary_data.append(
                    ["Total Candidates", str(report_data.candidate_summary.total_candidates)]
                )
                summary_data.append(
                    ["Verified", str(report_data.candidate_summary.verified_count)]
                )
                summary_data.append(
                    [
                        "Verification Rate",
                        f"{report_data.candidate_summary.verification_rate:.1%}",
                    ]
                )

            if summary_data:
                table = Table(summary_data, colWidths=[2 * inch, 2 * inch])
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("PADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                story.append(table)
            story.append(Spacer(1, 0.2 * inch))

        # Facilitator Performance
        if (
            ReportSection.FACILITATOR_PERFORMANCE in request.sections
            and report_data.facilitator_summaries
        ):
            story.append(Paragraph("Facilitator Performance", heading_style))

            fac_data = [["Facilitator", "Actions", "Candidates", "Verifications"]]
            for f in report_data.facilitator_summaries[:10]:
                fac_data.append(
                    [
                        f.user_name or f.user_id,
                        str(f.total_actions),
                        str(f.candidates_processed),
                        str(f.verification_actions),
                    ]
                )

            table = Table(fac_data, colWidths=[2 * inch, 1 * inch, 1.2 * inch, 1.2 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("PADDING", (0, 0), (-1, -1), 6),
                        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.2 * inch))

        # Conflict Resolution
        if ReportSection.CONFLICT_RESOLUTION in request.sections and report_data.conflict_summary:
            story.append(Paragraph("Conflict Resolution", heading_style))

            cf = report_data.conflict_summary
            conflict_data = [
                ["Total Conflicts", str(cf.total_conflicts)],
                ["Resolved", str(cf.resolved_conflicts)],
                ["Resolution Rate", f"{cf.resolution_rate:.1%}"],
            ]

            table = Table(conflict_data, colWidths=[2 * inch, 2 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("PADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.2 * inch))

        # Topic Trends
        if ReportSection.TOPIC_TRENDS in request.sections and report_data.topic_summaries:
            story.append(Paragraph("Top Topics", heading_style))

            topic_data = [["Topic", "Type", "Signals"]]
            for t in report_data.topic_summaries[:10]:
                topic_data.append([t.topic[:40], t.topic_type, str(t.signal_count)])

            table = Table(topic_data, colWidths=[3 * inch, 1.5 * inch, 1 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.darkblue),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("PADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.2 * inch))

        # Recommendations
        if ReportSection.RECOMMENDATIONS in request.sections and report_data.recommendations:
            story.append(Paragraph("Recommendations", heading_style))
            for rec in report_data.recommendations:
                story.append(Paragraph(f"• {rec}", body_style))
            story.append(Spacer(1, 0.2 * inch))

        # Build PDF
        doc.build(story)
        return buffer.getvalue()

    def _generate_docx(
        self,
        report_data: AfterActionReportData,
        request: AfterActionReportRequest,
    ) -> bytes:
        """Generate DOCX report using python-docx."""
        from docx import Document

        doc = Document()

        # Title
        doc.add_heading(report_data.title, 0)
        if report_data.incident_name:
            doc.add_paragraph(f"Incident: {report_data.incident_name}")
        doc.add_paragraph(
            f"Period: {report_data.start_date.strftime('%Y-%m-%d')} to "
            f"{report_data.end_date.strftime('%Y-%m-%d')}"
        )
        doc.add_paragraph(
            f"Generated: {report_data.generated_at.strftime('%Y-%m-%d %H:%M UTC')}"
        )

        # Executive Summary
        if ReportSection.EXECUTIVE_SUMMARY in request.sections:
            doc.add_heading("Executive Summary", level=1)

            if report_data.signal_summary:
                ss = report_data.signal_summary
                doc.add_paragraph(f"Total Signals: {ss.total_signals}")
                doc.add_paragraph(f"Average Signals per Day: {ss.avg_signals_per_day:.1f}")

            if report_data.candidate_summary:
                cs = report_data.candidate_summary
                doc.add_paragraph(f"Total Candidates: {cs.total_candidates}")
                doc.add_paragraph(f"Verified: {cs.verified_count}")
                doc.add_paragraph(f"Verification Rate: {cs.verification_rate:.1%}")

        # Facilitator Performance
        if (
            ReportSection.FACILITATOR_PERFORMANCE in request.sections
            and report_data.facilitator_summaries
        ):
            doc.add_heading("Facilitator Performance", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Facilitator"
            hdr_cells[1].text = "Actions"
            hdr_cells[2].text = "Candidates"
            hdr_cells[3].text = "Verifications"

            for f in report_data.facilitator_summaries[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = f.user_name or f.user_id
                row_cells[1].text = str(f.total_actions)
                row_cells[2].text = str(f.candidates_processed)
                row_cells[3].text = str(f.verification_actions)

        # Conflict Resolution
        if ReportSection.CONFLICT_RESOLUTION in request.sections and report_data.conflict_summary:
            doc.add_heading("Conflict Resolution", level=1)
            cf = report_data.conflict_summary
            doc.add_paragraph(f"Total Conflicts: {cf.total_conflicts}")
            doc.add_paragraph(f"Resolved: {cf.resolved_conflicts}")
            doc.add_paragraph(f"Resolution Rate: {cf.resolution_rate:.1%}")

        # Topic Trends
        if ReportSection.TOPIC_TRENDS in request.sections and report_data.topic_summaries:
            doc.add_heading("Top Topics", level=1)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Topic"
            hdr_cells[1].text = "Type"
            hdr_cells[2].text = "Signals"

            for t in report_data.topic_summaries[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = t.topic[:40]
                row_cells[1].text = t.topic_type
                row_cells[2].text = str(t.signal_count)

        # Recommendations
        if ReportSection.RECOMMENDATIONS in request.sections and report_data.recommendations:
            doc.add_heading("Recommendations", level=1)
            for rec in report_data.recommendations:
                doc.add_paragraph(f"• {rec}")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
